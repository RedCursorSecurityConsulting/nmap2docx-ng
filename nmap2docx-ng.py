import argparse
import os
import lxml.etree as ET
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.shared import OxmlElement, qn

def set_cell_background_color(cell, color, text_color=None):
    """ Set the background color of a cell. """
    shading_elm = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
    cell._tc.get_or_add_tcPr().append(shading_elm)
    if text_color:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor.from_string(text_color)

def parse_xml(xml_string):
    """ Parse XML for cell background color setting. """
    return ET.fromstring(xml_string)

def nsdecls(*prefixes):
    """ Generate namespace declarations for cell background color setting. """
    return ' '.join(['xmlns:%s="%s"' % (prefix, nsmap[prefix]) for prefix in prefixes])

nsmap = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
}

def parse_nmap_xml(xml_file_path):
    """ Parse the Nmap XML file. """
    tree = ET.parse(xml_file_path)
    root = tree.getroot()

    hosts_data = []
    for host in root.findall('.//host'):
        host_data = {'address': '', 'hostnames': '', 'ports': [], 'extraports': []}

        # Address
        for address in host.findall('address'):
            addr_type = address.get('addrtype')
            addr = address.get('addr')
            if addr_type == 'ipv4' or addr_type == 'ipv6':
                host_data['address'] += f"{addr} ({addr_type.upper()}), "

        host_data['address'] = host_data['address'].strip().strip(',')

        # Hostnames
        hostnames = host.find('hostnames')
        if hostnames is not None:
            hostnames_text = ', '.join([f"{hn.get('name')} ({hn.get('type')})" for hn in hostnames.findall('hostname')])
            host_data['hostnames'] = hostnames_text

        # Ports
        ports_element = host.find('ports')
        if ports_element is not None:
            for port in ports_element.findall('port'):
                port_data = {
                    'portid': port.get('portid'),
                    'protocol': port.get('protocol'),
                    'state': port.find('state').get('state'),
                    'service': port.find('service').get('name') if port.find('service') is not None else '', 
                    'product': port.find('service').get('product') if port.find('service') is not None else '',
                    'version': port.find('service').get('version') if port.find('service') is not None else ''
                }
                host_data['ports'].append(port_data)

            # Extra ports information
            extraports = ports_element.find('extraports')
            if extraports is not None:
                host_data['extraports'].append({
                    'state': extraports.get('state'),
                    'count': extraports.get('count')
                })

        # Adding host data to the list
        hosts_data.append(host_data)

    return hosts_data


def create_final_host_table(document, host_data):
    """ Create a table for each host in the Word document. """
    table = document.add_table(rows=0, cols=6)
    table.style = 'Table Grid'

    # Address row
    row = table.add_row()
    row.cells[0].merge(row.cells[5])
    row.cells[0].text = "Address"
    set_cell_background_color(row.cells[0], "E94347", "FFFFFF")

    row = table.add_row()
    row.cells[0].merge(row.cells[5])
    row.cells[0].text = host_data['address']

    # Hostnames row
    row = table.add_row()
    row.cells[0].merge(row.cells[5])
    row.cells[0].text = "Hostnames"
    set_cell_background_color(row.cells[0], "E94347", "FFFFFF")

    row = table.add_row()
    row.cells[0].merge(row.cells[5])
    row.cells[0].text = host_data['hostnames']

    # Column titles
    headers = ["Port", "Protocol", "State", "Service", "Product", "Version"]
    row = table.add_row()
    for i, title in enumerate(headers):
        row.cells[i].text = title
        set_cell_background_color(row.cells[i], "E94347", "FFFFFF")

    # Ports rows
    for port in host_data['ports']:
        row = table.add_row()
        row.cells[0].text = port['portid']
        row.cells[1].text = port['protocol']
        row.cells[2].text = port['state']
        row.cells[3].text = port['service']

        # Check if 'product' key exists in port_data, if not, set it to an empty string
        if 'product' in port and port['product'] is not None:
            row.cells[4].text = str(port['product'])
        else:
            row.cells[4].text = ''

        # Check if 'version' key exists in port_data, if not, set it to an empty string
        if 'version' in port and port['version'] is not None:
            row.cells[5].text = str(port['version'])
        else:
            row.cells[5].text = ''

        color = "EAF1DD" if port['state'] == 'open' else "F2DBDB"
        for cell in row.cells:
            set_cell_background_color(cell, color)

    # Extra ports info
    for extraport in host_data['extraports']:
        row = table.add_row()
        row.cells[0].merge(row.cells[5])
        row.cells[0].text = f"{extraport['count']} ports are in state: {extraport['state']}"

    # Set font for all cells
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Calibri'

    # Add a line break (paragraph) after each table
    document.add_paragraph()

# Make sure it's a valid NMAP XML file
def check_xml_file(xml_file_path):
    """ Check if the XML file is a valid Nmap file. """
    try:
        with open(xml_file_path, 'r') as file:
            first_line = file.readline()
            second_line = file.readline()
            if '<?xml version="1.0" encoding="UTF-8"?>' in first_line and '<!DOCTYPE nmaprun>' in second_line:
                return True
            else:
                return False
    except Exception as e:
        print(f"Error reading file: {e}")
        return False

def parse_arguments():
    """ Parse command line arguments. """
    parser = argparse.ArgumentParser(description='Process Nmap XML files.')
    parser.add_argument('-i', '--input', required=True, help='Input XML file path.')
    parser.add_argument('-o', '--output', help='Output DOCX file path (without .docx extension).', default='processed_hosts')
    args = parser.parse_args()
    return args.input, args.output

# Main script
input_file, output_file = parse_arguments()

if not check_xml_file(input_file):
    print("Invalid XML file.")
    print("Please generate a valid Nmap file with a command like:")
    print("sudo nmap -Pn -sC -sV -oX targets -iL targets.txt")
    exit(1)

hosts_data = parse_nmap_xml(input_file)

# Create a new document and add final tables for each host
doc = Document()
for host in hosts_data:
    create_final_host_table(doc, host)

# Handle output file naming
if not output_file.endswith('.docx'):
    output_file += '.docx'

# Save the final document
doc.save(output_file)
